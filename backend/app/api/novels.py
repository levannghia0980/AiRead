from fastapi import APIRouter, HTTPException, Depends, Query
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from typing import List, Optional
import io
import re
from urllib.parse import quote
import os
import hashlib
from sqlalchemy import select, delete, func
from sqlalchemy.ext.asyncio import AsyncSession
from app.core.database import get_db
from app.models.models import Novel, Chapter, Glossary, TranslationCache
from app.services.crawler.engine import scrape_novel_metadata
from app.services.exporter.packager import strip_html_tags

router = APIRouter(prefix="/api/novels", tags=["Novels"])

# Regex phát hiện chữ Hán còn sót trong bản dịch
_CHINESE_RE = re.compile(r"[\u4e00-\u9fff\u3400-\u4dbf\uf900-\ufaff]")

class AnalyzeRequest(BaseModel):
    url: str

class GlossaryCreate(BaseModel):
    chinese_term: str
    vietnamese_term: str
    category: str = "NAME"

class SaveChapterSchema(BaseModel):
    chapter_no: int
    title: str
    url: str

class SaveNovelRequest(BaseModel):
    title: str
    author: Optional[str] = "Khuyết Danh"
    cover_url: Optional[str] = ""
    source_url: str
    genres: Optional[str] = ""
    status: Optional[str] = "Ongoing"
    chapters: List[SaveChapterSchema]

class ResetChaptersRequest(BaseModel):
    chapter_nos: Optional[List[int]] = None

class UpdateChapterTextRequest(BaseModel):
    translated_text: str

@router.post("/analyze")
async def analyze_url(payload: AnalyzeRequest):
    """Scrapes novel meta details and lists of chapters from given URL."""
    if not payload.url:
        raise HTTPException(status_code=400, detail="Thiếu đường dẫn URL truyện")
    try:
        data = await scrape_novel_metadata(payload.url)
        return data
    except Exception as e:
        import traceback
        # Ghi log chi tiết lỗi ra file để debug
        log_file = r"d:\NENGHIA0980\AiRead2\error.log"
        try:
            with open(log_file, "w", encoding="utf-8") as f:
                traceback.print_exc(file=f)
        except Exception:
            pass
        raise HTTPException(status_code=500, detail=f"Lỗi phân tích URL: {str(e)}")

@router.get("")
async def list_novels(db: AsyncSession = Depends(get_db)):
    """Retrieves all registered novels from the database."""
    stmt = select(Novel).order_by(Novel.created_at.desc())
    result = await db.execute(stmt)
    return result.scalars().all()

@router.post("/save")
async def save_novel(payload: SaveNovelRequest, db: AsyncSession = Depends(get_db)):
    """Saves novel metadata and all chapters to database, facilitating resume operations."""
    # Check if already exists
    stmt = select(Novel).where(Novel.source_url == payload.source_url)
    result = await db.execute(stmt)
    existing_novel = result.scalar_one_or_none()
    
    if existing_novel:
        return {"novel_id": existing_novel.id, "message": "Truyện đã tồn tại sẵn trong hệ thống."}

    # Save new novel
    new_novel = Novel(
        title=payload.title,
        author=payload.author,
        cover_url=payload.cover_url,
        source_url=payload.source_url,
        genres=payload.genres,
        status=payload.status
    )
    db.add(new_novel)
    await db.flush() # populate ID

    # Batch save chapters
    db_chapters = [
        Chapter(
            novel_id=new_novel.id,
            chapter_no=ch.chapter_no,
            title=ch.title,
            source_url=ch.url,
            status="PENDING"
        )
        for ch in payload.chapters
    ]
    db.add_all(db_chapters)
    await db.commit()

    return {"novel_id": new_novel.id, "message": "Đã lưu truyện vào database thành công."}

@router.get("/{novel_id}")
async def get_novel_details(novel_id: int, db: AsyncSession = Depends(get_db)):
    """Gets novel metadata and its list of chapters."""
    novel = await db.get(Novel, novel_id)
    if not novel:
        raise HTTPException(status_code=444, detail="Không tìm thấy bộ truyện")
        
    stmt = select(Chapter).where(Chapter.novel_id == novel_id).order_by(Chapter.chapter_no.asc())
    result = await db.execute(stmt)
    chapters = result.scalars().all()
    
    return {
        "novel": novel,
        "chapters": chapters
    }

@router.delete("/{novel_id}")
async def delete_novel(novel_id: int, db: AsyncSession = Depends(get_db)):
    """Deletes novel and all its related chapters/glossaries."""
    novel = await db.get(Novel, novel_id)
    if not novel:
        raise HTTPException(status_code=404, detail="Không tìm thấy bộ truyện")
        
    await db.delete(novel)
    await db.commit()
    return {"success": True, "message": f"Đã xóa thành công truyện {novel.title}"}

# Glossary endpoints
@router.get("/{novel_id}/glossary")
async def get_glossary(novel_id: int, db: AsyncSession = Depends(get_db)):
    """Gets all glossaries matching this novel (and global ones)."""
    # Fetch terms for this novel or global (novel_id is null)
    stmt = select(Glossary).where((Glossary.novel_id == novel_id) | (Glossary.novel_id == None))
    result = await db.execute(stmt)
    return result.scalars().all()

@router.post("/{novel_id}/glossary")
async def add_glossary_term(novel_id: int, payload: GlossaryCreate, db: AsyncSession = Depends(get_db)):
    """Adds a glossary item to the novel (or global if novel_id is 0)."""
    target_novel_id = None if novel_id == 0 else novel_id
    
    # Check if duplicate exists
    stmt = select(Glossary).where(
        Glossary.novel_id == target_novel_id,
        Glossary.chinese_term == payload.chinese_term
    )
    result = await db.execute(stmt)
    existing = result.scalar_one_or_none()
    
    if existing:
        existing.vietnamese_term = payload.vietnamese_term
        existing.category = payload.category
    else:
        new_term = Glossary(
            novel_id=target_novel_id,
            chinese_term=payload.chinese_term,
            vietnamese_term=payload.vietnamese_term,
            category=payload.category
        )
        db.add(new_term)
        
    await db.commit()
    return {"success": True}

@router.delete("/{novel_id}/glossary/{glossary_id}")
async def delete_glossary_term(novel_id: int, glossary_id: int, db: AsyncSession = Depends(get_db)):
    """Deletes glossary item by ID."""
    stmt = delete(Glossary).where(Glossary.id == glossary_id)
    await db.execute(stmt)
    await db.commit()
    return {"success": True}

@router.get("/{novel_id}/chapters/{chapter_no}/text")
async def get_chapter_text(novel_id: int, chapter_no: int, db: AsyncSession = Depends(get_db)):
    """Gets a single chapter's translated text for reading."""
    stmt = select(Chapter).where(
        Chapter.novel_id == novel_id,
        Chapter.chapter_no == chapter_no
    )
    result = await db.execute(stmt)
    chapter = result.scalar_one_or_none()
    if not chapter:
        raise HTTPException(status_code=404, detail="Không tìm thấy chương")
    return {
        "chapter_no": chapter.chapter_no,
        "title": chapter.title,
        "status": chapter.status,
        "translated_text": chapter.translated_text or "",
        "raw_text": chapter.raw_text or ""
    }

@router.put("/{novel_id}/chapters/{chapter_no}/text")
async def update_chapter_text(
    novel_id: int,
    chapter_no: int,
    payload: UpdateChapterTextRequest,
    db: AsyncSession = Depends(get_db)
):
    """Updates the translated text of a single chapter."""
    stmt = select(Chapter).where(
        Chapter.novel_id == novel_id,
        Chapter.chapter_no == chapter_no
    )
    result = await db.execute(stmt)
    chapter = result.scalar_one_or_none()
    if not chapter:
        raise HTTPException(status_code=404, detail="Không tìm thấy chương")
    
    chapter.translated_text = payload.translated_text
    
    # Cập nhật lại file text vật lý nếu tồn tại để đồng bộ hóa
    novel = await db.get(Novel, novel_id)
    if novel:
        invalid_chars = '<>:"/\\|?*\r\n\t'
        safe_title = "".join(c for c in novel.title if c not in invalid_chars).strip()
        safe_title = safe_title.replace("  ", " ")
        BASE_DIR = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
        novel_folder = os.path.join(BASE_DIR, "output", safe_title)
        
        ch_title = "".join(c for c in chapter.title if c not in invalid_chars).strip()
        filename = f"Chương {chapter.chapter_no:04d} - {ch_title}.txt"
        filepath = os.path.join(novel_folder, filename)
        if os.path.exists(filepath):
            try:
                # Ghi vào file text đã lọc sạch HTML
                clean_txt = strip_html_tags(payload.translated_text)
                with open(filepath, "w", encoding="utf-8") as f:
                    f.write(f"{chapter.title}\n")
                    f.write("=" * 40 + "\n\n")
                    f.write(clean_txt)
            except Exception:
                pass
                
    await db.commit()
    return {"success": True, "message": "Đã lưu bản dịch chỉnh sửa thành công."}

@router.post("/{novel_id}/save-to-folder")
async def save_chapters_to_folder(novel_id: int, db: AsyncSession = Depends(get_db)):
    """Saves all translated chapters as individual .txt files in a folder named after the novel."""
    novel = await db.get(Novel, novel_id)
    if not novel:
        raise HTTPException(status_code=404, detail="Không tìm thấy truyện")
    
    # Create folder with novel title (Unicode-safe)
    invalid_chars = '<>:"/\\|?*\r\n\t'
    safe_title = "".join(c for c in novel.title if c not in invalid_chars).strip()
    safe_title = safe_title.replace("  ", " ")
    
    # Use project-level output directory
    BASE_DIR = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
    novel_folder = os.path.join(BASE_DIR, "output", safe_title)
    os.makedirs(novel_folder, exist_ok=True)
    
    # Fetch completed chapters
    stmt = (
        select(Chapter)
        .where(Chapter.novel_id == novel_id, Chapter.status == "COMPLETED")
        .order_by(Chapter.chapter_no.asc())
    )
    result = await db.execute(stmt)
    chapters = list(result.scalars().all())
    
    if not chapters:
        raise HTTPException(status_code=400, detail="Chưa có chương nào được dịch xong.")
    
    saved_files = []
    for ch in chapters:
        # Create clean filename: "Chương 001 - Title.txt"
        ch_title = "".join(c for c in ch.title if c not in invalid_chars).strip()
        filename = f"Chương {ch.chapter_no:04d} - {ch_title}.txt"
        filepath = os.path.join(novel_folder, filename)
        
        with open(filepath, "w", encoding="utf-8") as f:
            f.write(f"{ch.title}\n")
            f.write("=" * 40 + "\n\n")
            f.write(strip_html_tags(ch.translated_text or ""))
        
        saved_files.append(filename)
    
    return {
        "success": True,
        "folder": safe_title,
        "folder_path": novel_folder,
        "total_files": len(saved_files),
        "files": saved_files
    }

@router.post("/{novel_id}/chapters/reset")
async def reset_chapters(novel_id: int, payload: ResetChaptersRequest, db: AsyncSession = Depends(get_db)):
    """Resets chapters back to PENDING status and clears text to allow rescraping/retranslation."""
    stmt = select(Chapter).where(Chapter.novel_id == novel_id)
    if payload.chapter_nos is not None:
        stmt = stmt.where(Chapter.chapter_no.in_(payload.chapter_nos))
        
    result = await db.execute(stmt)
    chapters = result.scalars().all()
    
    if not chapters:
        raise HTTPException(status_code=404, detail="Không tìm thấy chương để reset")
        
    for ch in chapters:
        # Clear Translation Cache for this chapter's chunks if raw_text exists
        if ch.raw_text:
            try:
                # 1. Tách và dọn dẹp các đoạn văn để tìm các chunk chính xác (12000 ký tự) giống logic dịch
                paragraphs = [p.strip() for p in ch.raw_text.split("\n") if p.strip()]
                chunks = []
                current_chunk = []
                current_len = 0
                for p in paragraphs:
                    p_len = len(p)
                    if current_len + p_len > 12000 and current_chunk:
                        if p.startswith('"') and current_chunk[-1].startswith('"') and current_len + p_len < 12000 * 1.2:
                            current_chunk.append(p)
                            current_len += p_len + 2
                        else:
                            chunks.append("\n\n".join(current_chunk))
                            current_chunk = [p]
                            current_len = p_len
                    else:
                        current_chunk.append(p)
                        current_len += p_len + 2
                if current_chunk:
                    chunks.append("\n\n".join(current_chunk))
                
                # Tính MD5 của từng chunk
                chunk_hashes = [hashlib.md5(c.encode("utf-8")).hexdigest() for c in chunks if c.strip()]
                
                # 2. Xóa cache theo MD5 key_hash
                if chunk_hashes:
                    await db.execute(
                        delete(TranslationCache).where(
                            TranslationCache.key_hash.in_(chunk_hashes)
                        )
                    )
                
                # 3. Fallback: Xóa bằng so khớp chuỗi con (instr) để đảm bảo sạch 100% các đoạn lẻ khác
                await db.execute(
                    delete(TranslationCache).where(
                        func.instr(ch.raw_text, TranslationCache.raw_text) > 0
                    )
                )
            except Exception as e:
                import logging
                logging.getLogger(__name__).warning(f"Lỗi khi xóa cache của chương {ch.chapter_no}: {e}")

        ch.status = "PENDING"
        ch.raw_text = None
        ch.translated_text = None
        ch.error_msg = None
        
    # Also delete physical files for these chapters if they exist in the output folder
    novel = await db.get(Novel, novel_id)
    if novel:
        invalid_chars = '<>:"/\\|?*\r\n\t'
        safe_title = "".join(c for c in novel.title if c not in invalid_chars).strip()
        safe_title = safe_title.replace("  ", " ")
        BASE_DIR = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
        novel_folder = os.path.join(BASE_DIR, "output", safe_title)
        
        for ch in chapters:
            ch_title = "".join(c for c in ch.title if c not in invalid_chars).strip()
            filename = f"Chương {ch.chapter_no:04d} - {ch_title}.txt"
            filepath = os.path.join(novel_folder, filename)
            if os.path.exists(filepath):
                try:
                    os.remove(filepath)
                except Exception:
                    pass
                    
    # Xóa khỏi danh sách bỏ qua lỗi chất lượng để có thể quét lại sau khi dịch lại
    try:
        from app.models.models import Setting
        import json
        key = f"ignored_quality_{novel_id}"
        stmt_setting = select(Setting).where(Setting.key == key)
        res_setting = await db.execute(stmt_setting)
        setting_entry = res_setting.scalar_one_or_none()
        if setting_entry:
            ignored_list = json.loads(setting_entry.value)
            reset_nos = [ch.chapter_no for ch in chapters]
            new_list = [c for c in ignored_list if c not in reset_nos]
            setting_entry.value = json.dumps(new_list)
    except Exception:
        pass

    await db.commit()
    return {"success": True, "message": f"Đã reset {len(chapters)} chương về trạng thái chờ dịch."}


# Regex phát hiện câu tiếng Anh lẫn trong bản dịch (≥ 4 từ tiếng Anh liên tiếp)
_ENGLISH_PHRASE_RE = re.compile(r"(?:[A-Z][a-z]+(?:\s+[A-Za-z]+){3,})")
# Regex phát hiện cụm từ viết hoa kiểu thuật ngữ tiếng Anh (VD: "Central Continent Crossing Robbery Arena")
_ENGLISH_TERM_RE = re.compile(r"(?:[A-Z][a-z]+(?:\s+[A-Z][a-z]+){2,})")

@router.get("/{novel_id}/check-quality")
async def check_translation_quality(novel_id: int, db: AsyncSession = Depends(get_db)):
    """
    Quét nhanh tất cả chương đã dịch và phát hiện các lỗi phổ biến:
    - Còn chữ Hán (chữ Trung) trong bản dịch
    - Bản dịch rỗng hoặc quá ngắn so với bản gốc
    - Chương FAILED chưa được dịch lại
    - Dòng lặp lại nhiều lần (watermark / lỗi lặp)
    - Lặp liên tiếp 2+ dòng giống nhau
    - Câu tiếng Anh lẫn vào (dấu hiệu AI dịch sai ngu)
    - Câu dịch dự phòng (fallback) do AI chặn
    
    Trả về danh sách chương bị lỗi cùng mô tả lỗi.
    """
    # Tải danh sách chương đã đánh dấu bỏ qua/xác nhận OK
    from app.models.models import Setting
    import json
    key = f"ignored_quality_{novel_id}"
    stmt_setting = select(Setting).where(Setting.key == key)
    res_setting = await db.execute(stmt_setting)
    setting_entry = res_setting.scalar_one_or_none()
    ignored_chapters = set()
    if setting_entry:
        try:
            ignored_chapters = set(json.loads(setting_entry.value))
        except Exception:
            pass

    stmt = (
        select(Chapter)
        .where(Chapter.novel_id == novel_id)
        .order_by(Chapter.chapter_no.asc())
    )
    result = await db.execute(stmt)
    chapters = result.scalars().all()

    bad_chapters = []

    for ch in chapters:
        # Bỏ qua nếu chương này đã được người dùng đánh dấu là OK/Bỏ qua
        if ch.chapter_no in ignored_chapters:
            continue

        issues = []

        # 1. Chương FAILED chưa được dịch lại
        if ch.status == "FAILED":
            issues.append(f"❌ Lỗi dịch: {ch.error_msg or 'Không rõ nguyên nhân'}")

        # 2. Chương COMPLETED nhưng bản dịch bị vấn đề
        if ch.status == "COMPLETED":
            translated = ch.translated_text or ""

            # 2a. Bản dịch rỗng
            if not translated.strip():
                issues.append("❌ Bản dịch rỗng (không có nội dung)")

            else:
                # Strip HTML tags cho phân tích nội dung thuần
                plain_text = strip_html_tags(translated)
                
                # 2b. Còn chữ Hán trong bản dịch
                chinese_matches = _CHINESE_RE.findall(plain_text)
                if chinese_matches:
                    sample = "".join(chinese_matches[:20])
                    issues.append(f"🈲 Còn {len(chinese_matches)} chữ Hán sót lại (VD: {sample})")

                # 2c. Tỷ lệ độ dài quá thấp (dịch thiếu nội dung)
                raw = ch.raw_text or ""
                if raw.strip() and len(raw) > 100:
                    ratio = len(plain_text) / len(raw)
                    if ratio < 0.3:
                        issues.append(f"📉 Dịch thiếu nội dung: bản dịch chỉ bằng {ratio:.0%} bản gốc")

                # 2d. Kiểm tra dòng lặp lại ≥ 3 lần (watermark/lỗi lặp)
                lines = [l.strip() for l in plain_text.split("\n") if l.strip() and len(l.strip()) > 8]
                line_counts = {}
                for l in lines:
                    if len(l) > 12:
                        line_counts[l] = line_counts.get(l, 0) + 1
                
                dup_lines = {l: count for l, count in line_counts.items() if count >= 3}
                if dup_lines:
                    worst = max(dup_lines, key=dup_lines.get)
                    worst_count = dup_lines[worst]
                    sample = worst[:50] + "..." if len(worst) > 50 else worst
                    issues.append(f"🔁 Lặp {len(dup_lines)} câu (tệ nhất lặp {worst_count} lần: \"{sample}\")")

                # 2e. Lặp liên tiếp: 2 dòng giống nhau liền kề
                consecutive_dup_count = 0
                for i in range(1, len(lines)):
                    if lines[i] == lines[i - 1] and len(lines[i]) > 15:
                        consecutive_dup_count += 1
                if consecutive_dup_count >= 3:
                    issues.append(f"🔁 Có {consecutive_dup_count} cặp dòng lặp liên tiếp (dịch bị lặp nội dung)")

                # 2f. Phát hiện câu tiếng Anh lẫn vào (dấu hiệu AI dịch sai ngu)
                english_phrases = _ENGLISH_PHRASE_RE.findall(plain_text)
                english_terms = _ENGLISH_TERM_RE.findall(plain_text)
                # Lọc bỏ các cụm phổ biến cho phép (tên riêng ngắn, abbreviation...)
                real_english = [p for p in english_phrases if len(p) > 15]
                real_terms = [t for t in english_terms if len(t) > 15]
                
                if len(real_english) >= 2 or len(real_terms) >= 3:
                    samples = (real_english + real_terms)[:3]
                    sample_str = " | ".join(f'"{s[:40]}"' for s in samples)
                    issues.append(f"🇬🇧 Dịch sai ngu: có {len(real_english) + len(real_terms)} cụm tiếng Anh lẫn vào ({sample_str})")

                # 2g. Kiểm tra câu dịch dự phòng (fallback-line/fallback-word)
                fallback_line_count = translated.count('class="fallback-line"')
                fallback_word_count = translated.count('class="fallback-word"')
                total_fallback = fallback_line_count + fallback_word_count
                if total_fallback > 0:
                    details = []
                    if fallback_line_count > 0:
                        details.append(f"{fallback_line_count} câu dịch bù")
                    if fallback_word_count > 0:
                        details.append(f"{fallback_word_count} từ dịch thô")
                    issues.append(f"⚠️ Có {', '.join(details)} — cần kiểm tra lại")

                # 2h. Phát hiện bản dịch "sai ngu": quá nhiều câu cực ngắn (< 5 ký tự) so với bản gốc
                if len(lines) > 10:
                    very_short_lines = [l for l in lines if 0 < len(l) < 5]
                    short_ratio = len(very_short_lines) / len(lines)
                    if short_ratio > 0.3:
                        issues.append(f"🤡 Dịch ngu: {len(very_short_lines)}/{len(lines)} dòng cực ngắn (<5 ký tự), nghi ngờ dịch bị lỗi")

        if issues:
            bad_chapters.append({
                "chapter_no": ch.chapter_no,
                "title": ch.title,
                "status": ch.status,
                "issues": issues,
            })

    return {
        "total_chapters": len(chapters),
        "bad_count": len(bad_chapters),
        "bad_chapters": bad_chapters,
    }


# ===========================================================================
# TẢI FILE GỘP (TXT HOẶC DOCX)
# ===========================================================================

def clean_chapter_for_bundle(text: str) -> str:
    """
    Dọn dẹp nội dung chương để khi gộp sách được mạch lạc:
    - Loại bỏ các dòng lặp tiêu đề Việt (vd: "Chương 145: ...")
    - Loại bỏ các dòng tiêu đề Trung (vd: "第145章 ...")
    - Loại bỏ các dòng kết thúc chương (vd: "(Hết chương)", "(Hết chương này)")
    - Loại bỏ các dòng phân cách thô (vd: "-------", "======")
    """
    if not text:
        return ""
    
    # Strip HTML tags
    plain_text = strip_html_tags(text)
    
    lines = plain_text.split("\n")
    cleaned_lines = []
    
    # Regex nhận diện
    vi_chapter_re = re.compile(r"^\s*chương\s+\d+[:\s-]", re.IGNORECASE)
    zh_chapter_re = re.compile(r"^\s*第\s*\d+\s*[章章节]\s*", re.IGNORECASE)
    end_chapter_re = re.compile(r"^\s*[\(\[（]?\s*hết\s+chương\s*.*?[\)\]）]?\s*$", re.IGNORECASE)
    divider_re = re.compile(r"^\s*[-=_*]{3,}\s*$")
    
    for line in lines:
        stripped = line.strip()
        if not stripped:
            cleaned_lines.append("")
            continue
            
        # Kiểm tra và bỏ qua các dòng rác đầu/cuối chương
        if vi_chapter_re.match(stripped):
            continue
        if zh_chapter_re.match(stripped):
            continue
        if end_chapter_re.match(stripped):
            continue
        if divider_re.match(stripped):
            continue
            
        cleaned_lines.append(line)
        
    result = "\n".join(cleaned_lines)
    result = re.sub(r"\n{3,}", "\n\n", result)
    return result.strip()


@router.get("/{novel_id}/download")
async def download_novel(
    novel_id: int,
    fmt: str = Query("txt", description="Định dạng xuất: 'txt' hoặc 'docx'"),
    db: AsyncSession = Depends(get_db)
):
    """
    Gộp toàn bộ chương đã dịch thành một file duy nhất.
    Mỗi chương bắt đầu bằng tên chương, nội dung tuần tự từ chương 1 đến hết.
    Hỗ trợ định dạng: txt, docx
    """
    novel = await db.get(Novel, novel_id)
    if not novel:
        raise HTTPException(status_code=404, detail="Không tìm thấy truyện")

    stmt = (
        select(Chapter)
        .where(Chapter.novel_id == novel_id, Chapter.status == "COMPLETED")
        .order_by(Chapter.chapter_no.asc())
    )
    result = await db.execute(stmt)
    chapters = list(result.scalars().all())

    if not chapters:
        raise HTTPException(status_code=400, detail="Chưa có chương nào được dịch xong.")

    # Tên file an toàn
    invalid_chars = '<>:"/\\|?*\r\n\t'
    safe_title = "".join(c for c in novel.title if c not in invalid_chars).strip()
    safe_title = safe_title.replace("  ", " ") or "novel"

    fmt = fmt.lower().strip()

    # ── TXT ──────────────────────────────────────────────────────────────────
    if fmt == "txt":
        lines = []
        lines.append(novel.title)
        lines.append(f"Tác giả: {novel.author or 'Khuyết Danh'}")
        lines.append("=" * 60)
        lines.append("")

        for ch in chapters:
            cleaned_text = clean_chapter_for_bundle(ch.translated_text or "")
            if cleaned_text.strip():
                lines.append(cleaned_text)
                lines.append("")
                lines.append("")

        content = "\n".join(lines)
        buf = io.BytesIO(content.encode("utf-8"))
        buf.seek(0)
        filename = f"{safe_title}.txt"
        # RFC 5987: fallback ASCII + UTF-8 encoded filename for non-ASCII chars
        ascii_fallback = "novel.txt"
        encoded_name = quote(filename)

        return StreamingResponse(
            buf,
            media_type="text/plain; charset=utf-8",
            headers={"Content-Disposition": f"attachment; filename=\"{ascii_fallback}\"; filename*=UTF-8''{encoded_name}"}
        )

    # ── DOCX ─────────────────────────────────────────────────────────────────
    elif fmt == "docx":
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            raise HTTPException(
                status_code=500,
                detail="Thư viện python-docx chưa được cài. Chạy: pip install python-docx"
            )

        doc = Document()

        # Tiêu đề truyện
        title_para = doc.add_heading(novel.title, level=0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        author_para = doc.add_paragraph(f"Tác giả: {novel.author or 'Khuyết Danh'}")
        author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph("")

        for ch in chapters:
            # Nội dung chương — chia theo dòng để giữ định dạng đoạn sau khi lọc bỏ đầu cuối rác
            cleaned_text = clean_chapter_for_bundle(ch.translated_text or "")
            for para_text in cleaned_text.split("\n"):
                if para_text.strip():
                    doc.add_paragraph(para_text.strip())

            doc.add_paragraph("")  # Khoảng cách giữa các chương

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        filename = f"{safe_title}.docx"
        ascii_fallback = "novel.docx"
        encoded_name = quote(filename)

        return StreamingResponse(
            buf,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename=\"{ascii_fallback}\"; filename*=UTF-8''{encoded_name}"}
        )

    else:
        raise HTTPException(status_code=400, detail=f"Định dạng không hỗ trợ: {fmt}. Dùng 'txt' hoặc 'docx'.")


class QuickFixRequest(BaseModel):
    provider: str
    model: str
    api_key: str
    prompt: Optional[str] = None


@router.post("/{novel_id}/chapters/{chapter_no}/quick-fix")
async def quick_fix_chapter(
    novel_id: int,
    chapter_no: int,
    payload: QuickFixRequest,
    db: AsyncSession = Depends(get_db)
):
    """
    Sửa nhanh chương bị lỗi chất lượng (ví dụ: chứa câu tiếng Anh/pinyin,
    tên nhân vật dịch thô 'Lu Daguang' -> 'Lục Đại Hữu', câu chống cào...).
    
    Gửi ngữ cảnh đầy đủ của chương gồm bản gốc tiếng Trung và bản dịch thô
    tiếng Việt cho AI chỉnh sửa & biên tập lại chính xác.
    """
    import logging
    logger = logging.getLogger(__name__)

    # 1. Tìm chương truyện trong database
    stmt = select(Chapter).where(Chapter.novel_id == novel_id, Chapter.chapter_no == chapter_no)
    res = await db.execute(stmt)
    chapter = res.scalar_one_or_none()
    if not chapter:
        raise HTTPException(status_code=404, detail="Không tìm thấy chương")
    
    if not chapter.raw_text or not chapter.translated_text:
        raise HTTPException(status_code=400, detail="Chương chưa có đầy đủ nội dung để sửa nhanh")
    
    # 2. Khởi tạo AI Client
    from app.services.translator.client import TranslatorClient
    from app.services.translator.text_processor import postprocess_translated_text
    
    try:
        client = TranslatorClient(
            provider=payload.provider,
            model=payload.model,
            api_keys_str=payload.api_key,
            concurrency=1
        )
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Khởi tạo AI Client thất bại: {e}")
        
    # 3. Xây dựng System Instruction & Prompt sửa lỗi chuyên dụng
    system_instruction = (
        "Bạn là một biên tập viên dịch thuật tiểu thuyết Trung - Việt xuất sắc.\n"
        "Dưới đây là bản gốc tiếng Trung và bản dịch tiếng Việt hiện tại của một chương truyện.\n"
        "Bản dịch tiếng Việt hiện tại đang bị lỗi dịch máy thô, chứa cụm từ tiếng Anh vô nghĩa hoặc tên nhân vật bị viết sai kiểu pinyin/tiếng Anh (ví dụ: 'Lu Daguang' hoặc 'Lu Dayou' dịch đúng phải là 'Lục Đại Hữu', 'crossing robbery' -> 'độ kiếp' / 'vượt kiếp'). "
        "Nó cũng chứa các câu watermark chống cào của trang web Trung Quốc chèn rải rác.\n\n"
        "Nhiệm vụ của bạn:\n"
        "1. Quét qua bản dịch tiếng Việt và đối chiếu với bản gốc tiếng Trung để tìm các đoạn dịch lỗi/dịch thô/tối nghĩa.\n"
        "2. Sửa lại các đoạn dịch lỗi này sang văn phong tiếng Việt tiểu thuyết kiếm hiệp/tiên hiệp mượt mà, thuần Việt (ví dụ: sửa đúng tên nhân vật Lục Đại Hữu, Hàn Dục, độ kiếp...).\n"
        "3. Phát hiện và XÓA BỎ hoàn toàn các câu watermark chống cào quảng cáo nếu có.\n"
        "4. Giữ nguyên toàn bộ các câu/đoạn dịch chính xác khác của chương truyện.\n"
        "5. Trả về bản dịch tiếng Việt hoàn chỉnh sau khi đã sửa lỗi. Tuyệt đối không thêm bất kỳ lời giới thiệu, giải thích hay định dạng thừa nào ngoài bản dịch."
    )
    if payload.prompt:
        system_instruction += f"\n\nYêu cầu đặc biệt bổ sung từ người dùng: {payload.prompt}"
        
    plain_translated = strip_html_tags(chapter.translated_text)
    
    user_prompt = (
        f"BẢN GỐC TIẾNG TRUNG:\n{chapter.raw_text}\n\n"
        f"BẢN DỊCH TIẾNG VIỆT HIỆN TẠI (CÓ LỖI DỊCH MÁY/WATERMARK):\n{plain_translated}\n"
    )
    
    # 4. Gửi request cho AI sửa lỗi với cơ chế tự động giải cứu khi bị chặn (safety bypass)
    corrected_text = ""
    try:
        ai_res = await client.translate(user_prompt, system_instruction)
        corrected_text = ai_res.get("text", "").strip()
        if not corrected_text:
            raise Exception("AI returned empty corrected text")
    except Exception as first_err:
        logger.warning(f"⚠️ Sửa nhanh lần 1 bị AI chặn: {first_err}. Thử lại không dùng System Instruction...")
        try:
            # Thử lại không dùng system_instruction (hầu hết các bộ lọc nhạy cảm của Gemini sẽ giảm bớt)
            ai_res = await client.translate(user_prompt + "\n\n" + system_instruction, "")
            corrected_text = ai_res.get("text", "").strip()
            if not corrected_text:
                raise Exception("AI returned empty on attempt 2")
        except Exception as second_err:
            logger.warning(f"⚠️ Sửa nhanh lần 2 vẫn bị AI chặn: {second_err}. Tiến hành lách chữ nhạy cảm (Censor Bypass)...")
            try:
                # Thay thế các từ nhạy cảm dễ kích hoạt bộ lọc trong prompt bằng các ký tự đặc biệt
                safe_user_prompt = user_prompt
                # Các cụm từ nhạy cảm tiếng Việt
                sensitive_pairs = [
                    ("thuốc lắc", "thuốc l.ắ.c"),
                    ("người lớn mới được chơi", "người l.ớ.n ch.ơ.i"),
                    ("chết", "ch.ế.t"),
                    ("giết", "gi.ế.t"),
                    ("quan hệ", "q.u.a.n hệ"),
                    ("sát nhân", "s.á.t nhân"),
                    ("ác ý", "á.c ý"),
                    ("thắt lưng buộc quần", "thắt lưng"),
                ]
                for original, masked in sensitive_pairs:
                    safe_user_prompt = safe_user_prompt.replace(original, masked)
                
                # Gọi lại AI với prompt đã được làm sạch
                ai_res = await client.translate(safe_user_prompt, "")
                corrected_text = ai_res.get("text", "").strip()
                if not corrected_text:
                    raise Exception("AI returned empty on attempt 3")
                
                # Phục hồi lại các cụm từ gốc sau khi nhận kết quả
                for original, masked in sensitive_pairs:
                    corrected_text = corrected_text.replace(masked, original)
                        
            except Exception as third_err:
                raise HTTPException(
                    status_code=500,
                    detail=(
                        f"AI từ chối sửa lỗi chương này do vi phạm bộ lọc an toàn của Gemini (PROHIBITED_CONTENT).\n"
                        f"Chi tiết lỗi: {third_err}\n"
                        f"Mẹo: Bạn có thể cấu hình chọn nhà cung cấp AI khác (ví dụ: OpenRouter / DeepSeek-V3) ở cột bên phải để sửa nhanh không bị chặn."
                    )
                )
        
    # 5. Hậu xử lý (Enforce Glossary, dấu câu, khoảng trắng)
    from app.services.translator.pipeline import build_glossary_context
    stmt_glossary = select(Glossary).where(
        (Glossary.novel_id == 0) | (Glossary.novel_id == novel_id)
    )
    res_g = await db.execute(stmt_glossary)
    glossaries = res_g.scalars().all()
    
    _, glossary_map = build_glossary_context(chapter.raw_text, glossaries)
    
    final_text = await postprocess_translated_text(
        corrected_text,
        glossary_map,
        raw_chinese=chapter.raw_text,
        client=client
    )
    
    # 6. So sánh đối chiếu để tô màu xanh lá (text-cyber-success) các đoạn đã được sửa
    old_paras = [strip_html_tags(p).strip() for p in chapter.translated_text.split("\n") if p.strip()]
    new_paras = [p.strip() for p in final_text.split("\n") if p.strip()]
    
    old_set = set(old_paras)
    marked_paras = []
    for np in new_paras:
        # Nếu đoạn mới không khớp hoàn toàn với bất kỳ đoạn cũ nào, tức là đã được AI sửa đổi/dịch lại
        if np not in old_set:
            marked_paras.append(
                f'<span class="text-cyber-success bg-cyber-success/5 px-1.5 py-0.5 rounded border border-cyber-success/15 block my-1 font-medium">{np}</span>'
            )
        else:
            marked_paras.append(np)
            
    chapter.translated_text = "\n".join(marked_paras)
    
    # 7. Ghi đè file vật lý để đồng bộ
    novel = await db.get(Novel, novel_id)
    if novel:
        invalid_chars = '<>:"/\\|?*\r\n\t'
        safe_title = "".join(c for c in novel.title if c not in invalid_chars).strip()
        safe_title = safe_title.replace("  ", " ")
        BASE_DIR = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
        novel_folder = os.path.join(BASE_DIR, "output", safe_title)
        os.makedirs(novel_folder, exist_ok=True)
        
        ch_title = "".join(c for c in chapter.title if c not in invalid_chars).strip()
        filename = f"Chương {chapter.chapter_no:04d} - {ch_title}.txt"
        filepath = os.path.join(novel_folder, filename)
        try:
            clean_txt = strip_html_tags(final_text)
            with open(filepath, "w", encoding="utf-8") as f:
                f.write(f"{chapter.title}\n")
                f.write("=" * 40 + "\n\n")
                f.write(clean_txt)
        except Exception as e:
            logger.warning(f"Lỗi ghi file output khi sửa nhanh: {e}")
            
    # Tự động thêm vào danh sách bỏ qua quét lỗi (xác nhận chất lượng đã OK)
    try:
        from app.models.models import Setting
        import json
        key = f"ignored_quality_{novel_id}"
        stmt_setting = select(Setting).where(Setting.key == key)
        res_setting = await db.execute(stmt_setting)
        setting_entry = res_setting.scalar_one_or_none()
        if setting_entry:
            ignored_list = json.loads(setting_entry.value)
        else:
            setting_entry = Setting(key=key, value="[]")
            db.add(setting_entry)
            ignored_list = []
            
        if chapter_no not in ignored_list:
            ignored_list.append(chapter_no)
            setting_entry.value = json.dumps(ignored_list)
    except Exception as e:
        logger.warning(f"Lỗi tự động bỏ qua chất lượng sau sửa nhanh: {e}")

    await db.commit()
    return {"success": True, "message": f"Chương {chapter_no} đã được sửa nhanh bằng AI thành công!"}


@router.post("/{novel_id}/chapters/{chapter_no}/ignore-quality")
async def ignore_quality_chapter(
    novel_id: int,
    chapter_no: int,
    db: AsyncSession = Depends(get_db)
):
    """
    Đánh dấu một chương là 'Bỏ qua lỗi / Xác nhận OK'.
    Chương này sẽ không hiển thị trong danh sách quét lỗi chất lượng nữa.
    """
    from app.models.models import Setting
    import json
    
    key = f"ignored_quality_{novel_id}"
    stmt = select(Setting).where(Setting.key == key)
    res = await db.execute(stmt)
    setting = res.scalar_one_or_none()
    
    if setting:
        try:
            ignored_list = json.loads(setting.value)
        except Exception:
            ignored_list = []
    else:
        setting = Setting(key=key, value="[]")
        db.add(setting)
        ignored_list = []
        
    if chapter_no not in ignored_list:
        ignored_list.append(chapter_no)
        setting.value = json.dumps(ignored_list)
        
    await db.commit()
    return {"success": True, "message": f"Đã đánh dấu xác nhận OK (bỏ qua quét lỗi) cho chương {chapter_no}."}


@router.post("/{novel_id}/clear-ignored-quality")
async def clear_ignored_quality(
    novel_id: int,
    db: AsyncSession = Depends(get_db)
):
    """
    Xóa toàn bộ danh sách chương bỏ qua lỗi chất lượng của truyện này
    để có thể quét lỗi lại từ đầu cho toàn bộ truyện.
    """
    from app.models.models import Setting
    key = f"ignored_quality_{novel_id}"
    await db.execute(delete(Setting).where(Setting.key == key))
    await db.commit()
    return {"success": True, "message": "Đã reset danh sách bỏ qua lỗi chất lượng."}


