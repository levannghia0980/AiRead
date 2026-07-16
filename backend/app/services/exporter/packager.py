import os
import asyncio
from typing import List, Dict, Any
from sqlalchemy.future import select
from sqlalchemy.ext.asyncio import AsyncSession
from app.models.models import Novel, Chapter

# Conditional import helpers for docx and ebooklib to prevent crashes if they are not fully setup
import re

import html

def strip_html_tags(text: str) -> str:
    if not text:
        return ""
    # Remove HTML tags while preserving the inner content
    clean = re.sub(r'<[^>]+>', '', text)
    # Unescape HTML entities (e.g. &#20570; -> 做)
    clean = html.unescape(clean)
    return clean

def save_txt_file(filepath: str, title: str, author: str, chapters: List[Chapter]):
    with open(filepath, "w", encoding="utf-8") as f:
        f.write(f"TÊN TRUYỆN: {title}\n")
        f.write(f"TÁC GIẢ: {author}\n")
        f.write("=" * 40 + "\n\n")
        
        for ch in chapters:
            f.write(f"{ch.title}\n")
            f.write("-" * len(ch.title) + "\n\n")
            f.write(strip_html_tags(ch.translated_text or ""))
            f.write("\n\n" + "=" * 40 + "\n\n")

def clean_chapter_text(text: str, title: str) -> str:
    if not text:
        return ""
    
    lines = text.split("\n")
    cleaned_lines = []
    title_clean = title.strip().lower()
    skip_headers = True
    
    for line in lines:
        stripped = line.strip()
        if not stripped:
            if not skip_headers:
                cleaned_lines.append("")
            continue
            
        if skip_headers:
            # Check for matches with chapter title or common format patterns
            is_header = (
                stripped.lower() == title_clean or
                re.match(r"^(Chương|Chương thứ|Quyển|Tập)\s+\d+", stripped, re.IGNORECASE) or
                re.match(r"^第\s*\d+\s*(章|卷|回)", stripped) or
                stripped.lower() in title_clean or
                title_clean in stripped.lower() or
                stripped.startswith("=") or stripped.startswith("-") or stripped.startswith("*")
            )
            if is_header:
                continue
            else:
                skip_headers = False
                
        cleaned_lines.append(line)
        
    return "\n".join(cleaned_lines).strip()

def save_txt_clean_file(filepath: str, title: str, author: str, chapters: List[Chapter]):
    with open(filepath, "w", encoding="utf-8") as f:
        f.write(f"TÊN TRUYỆN: {title}\n")
        f.write(f"TÁC GIẢ: {author}\n")
        f.write("=" * 40 + "\n\n")
        
        for ch in chapters:
            cleaned_body = clean_chapter_text(ch.translated_text or "", ch.title)
            cleaned_body = strip_html_tags(cleaned_body)
            if cleaned_body:
                f.write(cleaned_body)
                f.write("\n\n")

def save_html_file(filepath: str, title: str, author: str, chapters: List[Chapter]):
    html_content = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="utf-8">
    <title>{title} - {author}</title>
    <style>
        body {{
            font-family: 'Palatino Linotype', 'Georgia', serif;
            line-height: 1.8;
            max-width: 800px;
            margin: 40px auto;
            padding: 0 20px;
            background-color: #fcfbf9;
            color: #2b2b2b;
        }}
        h1 {{
            text-align: center;
            font-size: 2.5em;
            color: #1a1a1a;
            margin-bottom: 5px;
        }}
        .author {{
            text-align: center;
            font-style: italic;
            color: #666;
            margin-bottom: 50px;
        }}
        h2 {{
            border-bottom: 1px solid #ddd;
            padding-bottom: 10px;
            margin-top: 60px;
            color: #333;
        }}
        .chapter-content {{
            font-size: 1.15em;
            text-align: justify;
            white-space: pre-line;
        }}
        .divider {{
            text-align: center;
            margin: 40px 0;
            color: #ccc;
            font-size: 1.5em;
        }}
        .fallback-word {
            color: #d97706;
            font-weight: bold;
            border-bottom: 1px dashed rgba(217, 119, 6, 0.6);
            background-color: rgba(217, 119, 6, 0.05);
        }
        .censor-word {
            color: #e11d48;
            font-weight: bold;
            border-bottom: 1px dashed rgba(225, 29, 72, 0.6);
            background-color: rgba(225, 29, 72, 0.05);
        }
        .fallback-line {
            border-bottom: 1px dotted rgba(0, 180, 216, 0.4);
            background-color: rgba(0, 180, 216, 0.02);
        }
        .censor-line {
            border-bottom: 1px dotted rgba(225, 29, 72, 0.4);
            background-color: rgba(225, 29, 72, 0.02);
        }
    </style>
</head>
<body>
    <h1>{title}</h1>
    <div class="author">Tác giả: {author}</div>
    <hr>
"""
    for ch in chapters:
        # Convert text body to paragraph tags
        text = ch.translated_text or ""
        paragraphs = text.split("\n\n")
        p_html = "".join(f"<p>{p.strip()}</p>" for p in paragraphs if p.strip())
        
        html_content += f"""
    <section class="chapter">
        <h2>{ch.title}</h2>
        <div class="chapter-content">
            {p_html}
        </div>
        <div class="divider">✦ ✦ ✦</div>
    </section>
"""
        
    html_content += """
</body>
</html>
"""
    with open(filepath, "w", encoding="utf-8") as f:
        f.write(html_content)

def save_docx_file(filepath: str, title: str, author: str, chapters: List[Chapter]):
    import docx
    doc = docx.Document()
    
    # Document title and info
    h = doc.add_heading(title, 0)
    h.alignment = 1 # Center
    
    p = doc.add_paragraph()
    p.add_run(f"Tác giả: {author}").italic = True
    p.alignment = 1 # Center
    
    doc.add_page_break()
    
    for ch in chapters:
        doc.add_heading(ch.title, level=1)
        text = strip_html_tags(ch.translated_text or "")
        paragraphs = text.split("\n\n")
        for para in paragraphs:
            if para.strip():
                doc.add_paragraph(para.strip())
        doc.add_paragraph("✦ ✦ ✦").alignment = 1 # Center
        
    doc.save(filepath)

def save_epub_file(filepath: str, title: str, author: str, chapters: List[Chapter]):
    from ebooklib import epub
    
    book = epub.EpubBook()
    book.set_identifier(f"airead_novel_{title.lower().replace(' ', '_')}")
    book.set_title(title)
    book.set_language("vi")
    book.add_author(author)
    
    epub_chapters = []
    
    # CSS style
    style = """
        @namespace epub "http://www.idpf.org/2007/ops";
        body {
            font-family: Georgia, serif;
            padding: 5%;
        }
        h2 {
            text-align: center;
            margin-bottom: 2em;
        }
        p {
            line-height: 1.6;
            text-indent: 1em;
            margin-bottom: 0.5em;
            text-align: justify;
        }
        .fallback-word {
            color: #d97706;
            font-weight: bold;
            border-bottom: 1px dashed rgba(217, 119, 6, 0.6);
        }
        .censor-word {
            color: #e11d48;
            font-weight: bold;
            border-bottom: 1px dashed rgba(225, 29, 72, 0.6);
        }
        .fallback-line {
            border-bottom: 1px dotted rgba(0, 180, 216, 0.4);
        }
        .censor-line {
            border-bottom: 1px dotted rgba(225, 29, 72, 0.4);
        }
    """
    default_css = epub.EpubItem(
        uid="style_default",
        file_name="style/default.css",
        media_type="text/css",
        content=style
    )
    book.add_item(default_css)
    
    for i, ch in enumerate(chapters, 1):
        ch_item = epub.EpubHtml(
            title=ch.title,
            file_name=f"chap_{i:04d}.xhtml",
            lang="vi"
        )
        
        # Convert text body to HTML paragraph tags
        text = ch.translated_text or ""
        paragraphs = text.split("\n\n")
        p_html = "".join(f"<p>{p.strip()}</p>" for p in paragraphs if p.strip())
        
        ch_item.content = f"<h2>{ch.title}</h2>\n{p_html}"
        ch_item.add_item(default_css)
        book.add_item(ch_item)
        epub_chapters.append(ch_item)
        
    book.toc = tuple(epub_chapters)
    
    # Required EPUB components
    book.add_item(epub.EpubNcx())
    book.add_item(epub.EpubNav())
    
    # Spine (reading order)
    book.spine = ["nav"] + epub_chapters
    
    epub.write_epub(filepath, book)


class NovelPackager:
    """Manages compilation and packaging of translated chapters into various book formats."""
    
    def __init__(self, output_dir: str):
        self.output_dir = output_dir
        os.makedirs(output_dir, exist_ok=True)

    async def package_novel(self, db: AsyncSession, novel_id: int) -> Dict[str, Any]:
        """
        Gathers all completed chapters and packs them into TXT, HTML, EPUB, and DOCX.
        Returns a dictionary of generated filenames.
        """
        # Fetch novel info
        stmt = select(Novel).where(Novel.id == novel_id)
        result = await db.execute(stmt)
        novel = result.scalar_one_or_none()
        
        if not novel:
            raise Exception(f"Novel not found with ID {novel_id}")
            
        # Fetch completed chapters in chronological order
        stmt = (
            select(Chapter)
            .where(Chapter.novel_id == novel_id)
            .where(Chapter.status == "COMPLETED")
            .order_by(Chapter.chapter_no.asc())
        )
        result = await db.execute(stmt)
        chapters = list(result.scalars().all())
        
        if not chapters:
            raise Exception("No completed chapters to package.")
            
        # Clean up book filenames — preserve Vietnamese/Unicode chars, only strip invalid file chars
        invalid_chars = '<>:"/\\|?*\r\n\t'
        safe_title = "".join(c for c in novel.title if c not in invalid_chars).strip()
        safe_title = safe_title.replace("  ", " ").replace(" ", "_")
        author = novel.author or "Khuyết Danh"
        
        txt_path = os.path.join(self.output_dir, f"{safe_title}.txt")
        txt_clean_path = os.path.join(self.output_dir, f"{safe_title}_lien_mach.txt")
        html_path = os.path.join(self.output_dir, f"{safe_title}.html")
        docx_path = os.path.join(self.output_dir, f"{safe_title}.docx")
        epub_path = os.path.join(self.output_dir, f"{safe_title}.epub")
        
        # Execute blocking I/O writes in thread pool
        loop = asyncio.get_running_loop()
        
        # Write TXT
        await loop.run_in_executor(None, save_txt_file, txt_path, novel.title, author, chapters)

        # Write Clean TXT (Liền mạch - bỏ tên chương)
        await loop.run_in_executor(None, save_txt_clean_file, txt_clean_path, novel.title, author, chapters)
        
        # Write HTML
        await loop.run_in_executor(None, save_html_file, html_path, novel.title, author, chapters)
        
        # Write DOCX (Try docx package)
        try:
            await loop.run_in_executor(None, save_docx_file, docx_path, novel.title, author, chapters)
            docx_ok = True
        except Exception as e:
            docx_ok = False
            print(f"Failed to export DOCX: {e}")
            
        # Write EPUB
        try:
            await loop.run_in_executor(None, save_epub_file, epub_path, novel.title, author, chapters)
            epub_ok = True
        except Exception as e:
            epub_ok = False
            print(f"Failed to export EPUB: {e}")
            
        return {
            "success": True,
            "title": novel.title,
            "txt": f"/output/{os.path.basename(txt_path)}",
            "txt_clean": f"/output/{os.path.basename(txt_clean_path)}",
            "html": f"/output/{os.path.basename(html_path)}",
            "docx": f"/output/{os.path.basename(docx_path)}" if docx_ok else None,
            "epub": f"/output/{os.path.basename(epub_path)}" if epub_ok else None,
        }
